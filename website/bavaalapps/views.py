import io
import os
from PIL import Image
from docx import Document
from docx.shared import Inches
from django.shortcuts import render
from django.http import HttpResponse, FileResponse
from django.conf import settings
from django.core.files.storage import FileSystemStorage
import json # Import json for handling session data
from .imagegridseo import pageseo 
 

# Helper function to create a blank white image placeholder
def create_blank_image(width, height):
    return Image.new('RGB', (width, height), (255, 255, 255))

# Helper function to create the Word document
def create_word(image_paths, texts, cols, target_size):
    doc = Document()

    # Calculate image width based on number of columns
    # 6.5 inches is approx page width with margins (standard A4/Letter)
    img_width_inches = 6.5 / cols
    img_width_emu = Inches(img_width_inches) # Convert to EMU for docx
    
    # Maintain aspect ratio for height
    if target_size[0] > 0: # Avoid division by zero
        img_height_inches = img_width_inches * (target_size[1] / target_size[0])
        img_height_emu = Inches(img_height_inches)
    else:
        img_height_emu = Inches(img_width_inches * 0.75) # Default aspect if width is zero

    # Ensure image_paths and texts are aligned
    num_items = max(len(image_paths), len(texts))
    
    for i in range(0, num_items, cols):
        # Get a row of images and texts
        row_image_paths = image_paths[i:i+cols]
        row_texts = texts[i:i+cols]

        # Pad the row if it doesn't have enough columns
        while len(row_image_paths) < cols:
            row_image_paths.append(None) # Use None for missing image paths
        while len(row_texts) < cols:
            row_texts.append("")

        # Create a table row
        table = doc.add_table(rows=1, cols=cols)
        table.autofit = False

        # Add images and text in row
        for j, (img_path, text) in enumerate(zip(row_image_paths, row_texts)):
            cell = table.cell(0, j)
            # Set cell width explicitly to match the calculated image width
            cell.width = img_width_emu

            if img_path:
                # Add image
                paragraph = cell.add_paragraph()
                run = paragraph.add_run()
                try:
                    run.add_picture(img_path, width=img_width_emu, height=img_height_emu)
                except Exception as e:
                    # Fallback for broken images or if path is invalid
                    paragraph.add_run(f"Error loading image: {e}")
                    print(f"Error adding picture {img_path}: {e}")
                
                # Add text below image
                cell.add_paragraph(text)
            else:
                # Add empty cell
                cell.text = ""
        # Add a paragraph after each table to ensure spacing between rows in the document
        doc.add_paragraph() 
    
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# Main view for the application
def index(request):
    # Clear session data on initial GET request to start fresh
    if request.method == 'GET':
        if 'uploaded_image_info' in request.session:
            del request.session['uploaded_image_info']
        if 'cols' in request.session:
            del request.session['cols']
        if 'target_width' in request.session:
            del request.session['target_width']
        if 'target_height' in request.session:
            del request.session['target_height']

    return render(request, 'bavaalapps/imagegrid.html', {'pageseo': pageseo})

# View to handle image uploads and display the grid
def upload_images(request):
    # Initialize context to pass to the template
    context = {
        'processed_images': [],
        'cols': 2, # Default columns
        'target_width': 400, # Default target width
        'target_height': 300, # Default target height
        'show_grid': False,
        'blank_image_url': '', # For blank placeholders
        'errors': [],
        'num_blank_placeholders': 0 # Initialize the new variable
    }

    if request.method == 'POST':
        # Get configuration from form
        try:
            cols = int(request.POST.get('cols', 2))
            target_width = int(request.POST.get('target_width', 400))
            target_height = int(request.POST.get('target_height', 300))
        except ValueError:
            context['errors'].append("Invalid number for columns, width, or height.")
            return render(request, 'bavaalapps/imagegrid.html', context)
        
        # Store configuration in session
        request.session['cols'] = cols
        request.session['target_width'] = target_width
        request.session['target_height'] = target_height

        uploaded_files = request.FILES.getlist('images')

        if not uploaded_files:
            context['errors'].append("Please upload at least one image.")
            context['cols'] = cols
            context['target_width'] = target_width
            context['target_height'] = target_height
            return render(request, 'bavaalapps/imagegrid.html', context)

        if len(uploaded_files) > 12:
            context['errors'].append("You can upload a maximum of 12 files.")
            return render(request, 'bavaalapps/imagegrid.html', context)

        # Define the temporary storage for this app only
        temp_media_root = os.path.join(settings.MEDIA_ROOT, 'tempfiles')
        temp_media_url = os.path.join(settings.MEDIA_URL, 'tempfiles')

         # Ensure the directory exists
        os.makedirs(temp_media_root, exist_ok=True)

        # Create a FileSystemStorage instance specific to this temporary location
        fs = FileSystemStorage(location=temp_media_root, base_url=temp_media_url)
        processed_image_info = [] # Store dicts with 'url' and 'path'

        for i, uploaded_file in enumerate(uploaded_files):
            try:
                # Save the uploaded file temporarily
                filename = fs.save(uploaded_file.name, uploaded_file)
                filepath = fs.path(filename)
                file_url = fs.url(filename)

                processed_image_info.append({
                    'url': file_url,
                    'path': filepath,
                    'caption': f"Image {i+1}" # Default caption
                })
            except Exception as e:
                context['errors'].append(f"Error processing {uploaded_file.name}: {e}")
                continue

        # Calculate number of blank placeholders needed to complete the last row
        num_images = len(processed_image_info)
        num_blank_placeholders = 0
        if num_images > 0: # Only add placeholders if there are some images
            remainder = num_images % cols
            if remainder != 0:
                num_blank_placeholders = cols - remainder

        # Store processed image info in session
        request.session['uploaded_image_info'] = json.dumps(processed_image_info)
        
        context.update({
            'processed_images': processed_image_info,
            'cols': cols,
            'target_width': target_width,
            'target_height': target_height,
            'show_grid': True,
            'num_blank_placeholders': num_blank_placeholders # Pass the calculated value
        })

    return render(request, 'bavaalapps/imagegrid.html', context)

# View to generate and download the Word document
def generate_word_document(request):
    if request.method == 'POST':
        # Retrieve image info and configuration from session
        uploaded_image_info_json = request.session.get('uploaded_image_info')
        if not uploaded_image_info_json:
            return HttpResponse("No images found to generate document. Please upload images first.", status=400)
        
        uploaded_image_info = json.loads(uploaded_image_info_json)
        cols = request.session.get('cols', 2)
        target_width = request.session.get('target_width', 400)
        target_height = request.session.get('target_height', 300)

        image_paths = []
        texts = []

        # Collect captions and actual image paths
        for i, img_info in enumerate(uploaded_image_info):
            caption_key = f'caption_{i}'
            caption = request.POST.get(caption_key, img_info.get('caption', f"Image {i+1}"))
            texts.append(caption)
            image_paths.append(img_info['path'])
        
        # Generate the Word document
        try:
            doc_buffer = create_word(image_paths, texts, cols, (target_width, target_height))
        except Exception as e:
            return HttpResponse(f"Error generating Word document: {e}", status=500)

        response = FileResponse(
            doc_buffer,
            as_attachment=True,
            filename="bavaal_tools_grid.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        return response
    return HttpResponse("Invalid request.", status=400)